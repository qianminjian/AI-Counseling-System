#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph_with_style(doc, text, style_name='Normal'):
    """添加段落"""
    p = doc.add_paragraph(text, style=style_name)
    return p

def create_table_with_header(doc, headers, rows, col_widths=None):
    """创建带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'D9E2F3')

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)

    return table

def main():
    doc = Document()

    # 设置文档标题
    title = doc.add_heading('AI心理辅导系统商业模式与学校采购逻辑分析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ========== 一、市场分析 ==========
    add_heading(doc, '一、市场分析', 1)

    add_heading(doc, '1.1 K12教育市场规模', 2)
    p = doc.add_paragraph()
    p.add_run('中国K12教育市场规模持续增长，2024年市场规模约达5.2万亿元人民币。心理健康教育作为素质教育的重要组成部分，占据越来越重要的地位。随着政策推动和社会认知提升，校园心理健康服务市场呈现快速增长态势。')

    add_heading(doc, '1.2 心理健康教育政策驱动', 2)
    policies = [
        ('2019年', '国家卫健委等12部门发布《健康中国行动——儿童青少年心理健康行动方案（2019—2022年）》'),
        ('2021年', '教育部办公厅发布《关于加强学生心理健康管理工作的通知》，明确要求中小学配备心理辅导室'),
        ('2022年', '《中华人民共和国未成年人保护法》修订，明确学校保护责任'),
        ('2023年', '《全面加强和改进新时代学生心理健康工作专项行动计划（2023—2025年）》'),
        ('2024年', '《学生心理健康工作体系建设三年行动计划》启动实施'),
    ]
    doc.add_paragraph('近年来，国家层面出台多项政策推动校园心理健康工作：')
    for year, policy in policies:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{year}：{policy}')

    add_heading(doc, '1.3 学校心理辅导现状与痛点', 2)
    pain_points = [
        ('专业人才匮乏', '全国仅有约3万名专职中小学心理健康教育教师，师生比远低于发达国家水平'),
        ('筛查效率低下', '传统心理筛查依赖纸质问卷，人工分析耗时长、主观性强、覆盖面窄'),
        ('预警机制缺失', '缺乏实时监测和早期预警手段，往往在问题爆发后才被发现'),
        ('家校协同困难', '心理问题识别和干预需要多方协同，但缺乏有效的信息共享机制'),
        ('隐私保护不足', '学生心理数据的采集、存储、使用缺乏规范，存在隐私泄露风险'),
    ]
    for title, desc in pain_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}：').bold = True
        p.add_run(desc)

    add_heading(doc, '1.4 目标客户画像', 2)

    customers = [
        ['客户类型', '特征描述', '核心需求', '采购决策权'],
        ['重点中学', '高考压力大、心理问题突出', '高效筛查、快速预警', '校长、教务处'],
        ['普通中小学', '资源有限、人员不足', '性价比、简单易用', '校长、总务处'],
        ['教育局', '统筹区域学校管理', '统一平台、数据汇总', '局长、科室负责人'],
        ['国际学校', '家长期望高、注重隐私', '高品质服务、数据安全', '校董会、外籍校长'],
    ]
    create_table_with_header(doc, customers[0], customers[1:], [1.5, 3, 3, 1.5])

    doc.add_paragraph()

    # ========== 二、商业模式设计 ==========
    add_heading(doc, '二、商业模式设计', 1)

    add_heading(doc, '2.1 SaaS订阅模式', 2)

    p = doc.add_paragraph()
    p.add_run('采用SaaS订阅模式，降低学校前期投入成本，提高产品可及性。').bold = True

    sub_sections = [
        ('按学校/按学生收费', [
            '按学校收费：适合规模较大、预算独立的重点学校',
            '按学生收费：适合统一采购的教育局集采项目',
            '建议采用"基础学生数+超出部分按量计费"模式'
        ]),
        ('年费制vs月费制', [
            '年费制：给予8-9折优惠，降低续费摩擦，提高现金流可预测性',
            '月费制：适合试点项目或短期需求，灵活性高但单价较高',
            '建议主流推广年费制，季度付为辅'
        ]),
        ('分层定价（基础版/专业版/旗舰版）', [
            '基础版（99元/生/年）：核心筛查+基础报告+单校管理',
            '专业版（159元/生/年）：AI对话+危机预警+家校互通+教师培训',
            '旗舰版（259元/生/年）：全功能+专属实施+区域数据看板+优先支持'
        ]),
    ]

    for sub_title, items in sub_sections:
        add_heading(doc, sub_title, 3)
        for item in items:
            p = doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '2.2 采购决策链分析', 2)

    decision_chain = [
        ['角色', '职责', '关注点', '影响策略'],
        ['决策者\n（校长/教育局）', '最终批准采购', '政策合规、舆论风险、投入产出比', '政策解读、案例背书、风险兜底承诺'],
        ['影响者\n（心理老师/班主任）', '提供采购建议', '功能实用性、工作量降低', '产品培训、操作演示、满意度保障'],
        ['使用者\n（学生）', '实际使用产品', '体验流畅、隐私保护', '趣味设计、正向反馈机制'],
        ['评估者\n（家委会）', '监督服务质量', '安全性、效果可见', '透明报告、家长开放日'],
    ]
    create_table_with_header(doc, decision_chain[0], decision_chain[1:])

    doc.add_paragraph()

    add_heading(doc, '2.3 采购触发点', 2)

    triggers = [
        ('政策要求', '必须配备', '教育部/省厅明确要求，心理测评系统成为学校达标项', '教育局红头文件、达标验收时间节点'),
        ('事件驱动', '心理事件发生', '学生心理危机事件引起重视，主动寻求预防工具', '突发事件、媒体报道、同类学校案例'),
        ('指标考核', '心理筛查要求', '将心理健康筛查率、预警准确率纳入学校考核', '考核周期、考核指标、排名压力'),
        ('竞品对比', '替换现有方案', '现有系统功能落后或服务不佳，考虑更换', '竞品弱点、本产品差异化优势'),
    ]

    for trigger_type, condition, desc, timing in triggers:
        p = doc.add_paragraph()
        p.add_run(f'{trigger_type}：').bold = True
        p.add_run(f'{condition} — {desc}（{timing}）')

    # ========== 三、定价策略 ==========
    add_heading(doc, '三、定价策略', 1)

    add_heading(doc, '3.1 参考竞品定价', 2)
    competitors = [
        ['竞品', '定价区间', '定位', '备注'],
        ['心理云平台', '80-150元/生/年', '中端市场', '功能较全但AI能力弱'],
        ['测评公司传统软件', '50-100元/生/年', '中低端市场', '本地部署为主'],
        ['互联网大厂产品', '100-200元/生/年', '中高端市场', '品牌强但定制化弱'],
        ['国际心理测评系统', '300-500元/生/年', '高端市场', '国际学校为主'],
    ]
    create_table_with_header(doc, competitors[0], competitors[1:])
    doc.add_paragraph()

    add_heading(doc, '3.2 成本加成定价', 2)
    cost_items = [
        ('AI模型成本', '约15-20元/生/年（对话+分析）'),
        ('服务器与带宽', '约5-8元/生/年'),
        ('实施与支持', '约10-15元/生/年（首年较高）'),
        ('研发摊销', '约8-12元/生/年'),
        ('毛利率目标', '60-70%'),
    ]
    for item, desc in cost_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    add_heading(doc, '3.3 价值定价', 2)
    values = [
        ('政策合规价值', '满足心理测评系统配置要求，规避合规风险', '5-10元/生'),
        ('效率提升价值', 'AI自动分析替代人工，节省心理老师工作量', '10-20元/生'),
        ('预警避险价值', '早期发现心理问题，降低危机事件发生概率', '15-30元/生'),
        ('数据洞察价值', '区域心理健康数据看板，支撑管理决策', '10-15元/生'),
    ]
    for value, desc, price in values:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{value}：{desc}（{price}）')

    add_heading(doc, '3.4 分层功能定价表', 2)

    pricing_table = [
        ['功能模块', '基础版', '专业版', '旗舰版'],
        ['心理筛查', '✓', '✓', '✓'],
        ['AI智能对话', '-', '✓', '✓'],
        ['危机预警', '-', '✓', '✓'],
        ['家校互通', '-', '✓', '✓'],
        ['区域数据看板', '-', '-', '✓'],
        ['API开放', '-', '-', '✓'],
        ['专属实施', '-', '-', '✓'],
        ['SLA保障', '-', '99.5%', '99.9%'],
        ['定价（元/生/年）', '99', '159', '259'],
    ]
    create_table_with_header(doc, pricing_table[0], pricing_table[1:])

    # ========== 四、销售策略 ==========
    add_heading(doc, '四、销售策略', 1)

    add_heading(doc, '4.1 标杆学校打法', 2)
    steps = [
        ('选择标杆', '选取2-3所具有影响力的重点中学作为首批试点'),
        ('深度合作', '与学校心理老师共创优化产品，形成可复制的使用模式'),
        ('效果验证', '通过实际数据展示筛查效率提升、预警准确率等核心指标'),
        ('案例包装', '形成标准化的标杆案例材料，包含量化成果和用户证言'),
        ('复制推广', '以标杆学校为中心，向同区域其他学校辐射'),
    ]
    for i, (step, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {step}：').bold = True
        p.add_run(desc)

    add_heading(doc, '4.2 教育局集采路径', 2)
    p = doc.add_paragraph('适用于区域统一采购场景：')
    edu_steps = [
        ('切入点选择', '以"心理健康筛查覆盖率"考核指标为切入点'),
        ('试点先行', '选取1-2个区县作为试点，积累数据和口碑'),
        ('汇报路径', '通过教育局内部简报、现场会等形式触达决策层'),
        ('集采谈判', '争取纳入年度采购计划或教育信息化项目'),
        ('服务延伸', '从区县级向市级、省级推广'),
    ]
    for step, desc in edu_steps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{step}：').bold = True
        p.add_run(desc)

    add_heading(doc, '4.3 渠道合作（代理商/集成商）', 2)
    channels = [
        ('代理商模式', '发展本地教育信息化代理商，给予20-30%渠道利润空间'),
        ('集成商合作', '与智慧校园、校园一卡通等集成商合作，进入其产品矩阵'),
        ('战略合作', '与心理测评机构、公立青少年心理援助中心等建立合作'),
    ]
    for channel, desc in channels:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{channel}：').bold = True
        p.add_run(desc)

    add_heading(doc, '4.4 试用转化策略', 2)
    trial = [
        ('免费试用期', '提供14-30天全功能免费试用'),
        ('试用目标', '完成至少1次完整心理测评流程，体验核心价值'),
        ('转化节点', '试用期结束前3天跟进，收集反馈并促成签单'),
        ('优惠策略', '试用期结束后1周内签约，享受年费9折优惠'),
    ]
    for item, desc in trial:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    # ========== 五、合同条款建议 ==========
    add_heading(doc, '五、合同条款建议', 1)

    add_heading(doc, '5.1 数据归属条款', 2)
    data_terms = [
        ('数据所有权', '明确约定学校/教育局对学生心理数据拥有所有权'),
        ('数据存储', '约定数据存储地点、期限及销毁方式'),
        ('数据使用', '明确我方使用数据的范围（如匿名化统计、产品优化）'),
        ('数据迁移', '合同终止时支持数据导出，格式需双方确认'),
    ]
    for term, desc in data_terms:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{term}：').bold = True
        p.add_run(desc)

    add_heading(doc, '5.2 服务等级协议（SLA）', 2)
    sla_terms = [
        ('可用性承诺', '基础版99.5%、专业版99.7%、旗舰版99.9%'),
        ('响应时间', '普通问题4小时、紧急问题1小时、危机问题15分钟'),
        ('赔偿机制', '可用性每下降0.1%补偿1天服务时长'),
        ('不可抗力', '明确自然灾害、重大公共事件等免责情形'),
    ]
    for term, desc in sla_terms:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{term}：').bold = True
        p.add_run(desc)

    add_heading(doc, '5.3 保密条款', 2)
    privacy_terms = [
        ('保密义务', '双方对合作过程中获取的对方商业秘密负有保密义务'),
        ('保密期限', '合同终止后仍需保密不少于3年'),
        ('数据脱敏', '对外展示案例需经过脱敏处理并获得授权'),
    ]
    for term, desc in privacy_terms:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{term}：').bold = True
        p.add_run(desc)

    add_heading(doc, '5.4 退出条款', 2)
    exit_terms = [
        ('提前通知期', '任一方提前60天书面通知可解除合同'),
        ('数据交接', '合同终止后30天内完成全部数据交接'),
        ('费用结算', '按实际使用时长结算，不收取违约金'),
        ('竞业限制', '可选条款，限制我方在特定区域开展同类业务'),
    ]
    for term, desc in exit_terms:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{term}：').bold = True
        p.add_run(desc)

    # ========== 六、客户成功策略 ==========
    add_heading(doc, '六、客户成功策略', 1)

    add_heading(doc, '6.1 实施交付标准', 2)
    implementation = [
        ('项目启动', '明确负责人、对接人、项目目标'),
        ('数据准备', '协助学校完成学生信息导入、账号开通'),
        ('系统配置', '完成校情定制、筛查量表配置'),
        ('培训实施', '完成教师端、学生端操作培训'),
        ('试运行', '2周试运行，收集反馈并优化'),
        ('正式上线', '正式投入使用，移交文档资料'),
    ]
    for step, desc in implementation:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{step}：').bold = True
        p.add_run(desc)

    add_heading(doc, '6.2 教师培训方案', 2)
    training = [
        ('新手培训', '线上视频课程（约2小时）+操作手册'),
        ('进阶培训', '月度线上答疑会，解答实际使用问题'),
        ('专项培训', '针对心理老师的危机识别与干预培训'),
        ('培训考核', '培训后进行简单考核，确保操作熟练度'),
    ]
    for item, desc in training:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    add_heading(doc, '6.3 持续运营支持', 2)
    support = [
        ('专属客服', '企业微信群+电话支持，4小时响应'),
        ('定期回访', '每月1次使用情况回访，收集改进建议'),
        ('版本迭代', '每季度发布新功能，持续优化体验'),
        ('活动支持', '配合学校心理健康教育月等活动提供素材'),
    ]
    for item, desc in support:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    add_heading(doc, '6.4 续费激励机制', 2)
    renewal = [
        ('早鸟优惠', '合同到期前60天续费享9折'),
        ('忠诚折扣', '连续3年续费享8.5折'),
        ('增值赠送', '续费赠送额外学生账号或功能模块'),
        ('案例激励', '愿意共享案例的客户可获得额外服务时长'),
    ]
    for item, desc in renewal:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    # ========== 七、竞争壁垒 ==========
    add_heading(doc, '七、竞争壁垒', 1)

    add_heading(doc, '7.1 心理学专业壁垒', 2)
    p = doc.add_paragraph('心理健康领域对专业性要求极高，我方需要建立以下专业壁垒：')
    professional = [
        ('量表研发', '拥有自主知识产权的本土化心理评估量表'),
        ('AI模型训练', '基于大量真实案例训练的AI对话和预警模型'),
        ('专家团队', '聘请资深心理咨询师担任顾问，参与产品设计审核'),
        ('学术合作', '与高校心理学院建立产学研合作'),
    ]
    for item, desc in professional:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    add_heading(doc, '7.2 学校关系壁垒', 2)
    school_relation = [
        ('用户习惯', '师生已习惯现有操作流程，更换成本高'),
        ('数据积累', '长期使用形成的历史数据资产'),
        ('信任关系', '通过多次成功服务建立的信任关系'),
        ('区域网络', '在特定区域形成的良好口碑和网络效应'),
    ]
    for item, desc in school_relation:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    add_heading(doc, '7.3 数据壁垒', 2)
    data_barrier = [
        ('数据规模', '覆盖学生越多，AI模型越精准，形成正向循环'),
        ('数据质量', '经过标注和验证的高质量训练数据'),
        ('数据隐私', '在合规框架下探索数据价值最大化'),
    ]
    for item, desc in data_barrier:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    add_heading(doc, '7.4 政策合规壁垒', 2)
    policy_barrier = [
        ('资质认证', '通过教育部教育APP备案、等级保护认证'),
        ('隐私合规', '符合《个人信息保护法》《未成年人保护法》要求'),
        ('行业标准', '参与行业标准制定，抢占先发优势'),
    ]
    for item, desc in policy_barrier:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    # ========== 八、风险与应对 ==========
    add_heading(doc, '八、风险与应对', 1)

    risks = [
        ['风险类型', '具体风险', '发生概率', '影响程度', '应对策略'],
        ['政策风险', '政策收紧或标准变化', '中', '高', '密切跟踪政策动向，保持产品灵活性，参与标准制定'],
        ['竞争风险', '大厂低价竞争', '高', '中', '深耕专业能力，强化服务差异化，聚焦细分市场'],
        ['交付风险', '实施能力跟不上订单增长', '中', '中', '建立标准化交付体系，培养合作实施伙伴'],
        ['现金流风险', '回款周期长或坏账', '中', '高', '优化合同条款，加强应收账款管理，储备现金'],
        ['数据安全风险', '数据泄露或被攻击', '低', '极高', '加强安全投入，定期渗透测试，购买网络安全险'],
        ['人才流失风险', '核心人员离职', '中', '中', '建立知识库，培养梯队，完善激励机制'],
    ]
    create_table_with_header(doc, risks[0], risks[1:])

    # ========== 九、三年财务预测 ==========
    add_heading(doc, '九、三年财务预测（简表）', 1)

    add_heading(doc, '9.1 核心假设', 2)
    assumptions = [
        ('客户增长', '第1年50所学校，第2年150所，第3年350所'),
        ('平均客单价', '第1年120元/生（早期优惠），第2年135元/生，第3年145元/生'),
        ('平均学校规模', '1000名学生/校'),
        ('续费率', '第2年80%，第3年85%'),
        ('毛利率', '65%（稳定期）'),
    ]
    for item, desc in assumptions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}：').bold = True
        p.add_run(desc)

    add_heading(doc, '9.2 收入预测', 2)

    financial_table = [
        ['项目', '第1年', '第2年', '第3年'],
        ['累计学校数', '50', '200', '550'],
        ['当年新增学校', '50', '150', '350'],
        ['当年续费学校', '0', '40', '160'],
        ['活跃学生数（万）', '5', '20', '55'],
        ['平均客单价（元/生）', '120', '135', '145'],
        ['年度总收入（万元）', '600', '2,700', '7,975'],
        ['同比增长率', '-', '350%', '195%'],
        ['主营业务成本（万元）', '210', '945', '2,791'],
        ['毛利（万元）', '390', '1,755', '5,184'],
        ['毛利率', '65%', '65%', '65%'],
        ['运营费用（万元）', '800', '1,600', '2,800'],
        ['净利润（万元）', '-410', '155', '2,384'],
        ['累计净利润（万元）', '-410', '-255', '2,129'],
    ]
    create_table_with_header(doc, financial_table[0], financial_table[1:])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('说明：').bold = True
    p.add_run('第1年处于市场拓展期，收入较低且需投入大量市场和销售费用，预计亏损410万元。第2年随着口碑建立和渠道拓展，实现盈亏平衡并小幅盈利。第3年在续费收入稳定增长和新客户持续拓展下，实现规模化盈利。')

    add_heading(doc, '9.3 关键指标监控', 2)
    kpis = [
        ('获客成本（CAC）', '<5,000元/校'),
        ('客户生命周期价值（LTV）', '>30,000元/校'),
        ('LTV/CAC比值', '>3'),
        ('续费率', '>80%'),
        ('净推荐值（NPS）', '>50'),
        ('客户回收期', '<18个月'),
    ]
    for kpi, target in kpis:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{kpi}：').bold = True
        p.add_run(target)

    # 保存文档
    output_path = '/Users/minjianq/Documents/AI-Counseling-System/PRD/子主题/09_商业模式与学校采购逻辑.docx'
    doc.save(output_path)
    print(f'文档已保存至: {output_path}')

if __name__ == '__main__':
    main()
